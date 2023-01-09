import datetime
import json
import os.path
import random
import re
import shutil
import sys
import tarfile
import urllib.error
import urllib.request
from collections import OrderedDict
from pathlib import Path
import docx
import docx2txt
import unidecode
import yaml
from docx.document import Document as doc
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

config = {
    "directories": {
        "input": "source",
        "output": "out",
        "themes": "themes",
    },
    "site": {
        "theme": "default",
        "domain": "wordflow.com"
    },
    "author": {
        "nickname": "ahr",
        "name": "author",
        "email": "you@me.com",
        "about": "I publish my word documents using wordflow!",
    },
    "generator": {
        "input": "docx",  # input language (md soon..)
    }
}
content = {}

theme = {}

styles = {
    "Title": "h1",
    "Heading 1": "h1",
    "Heading 2": "h2",
    "Heading 3": "h3",
    "Emphasis": "u",
    "Normal": "p",
    "List Paragraph": "li",
    "List Number": "li",
    "List Bullet": "li",
    "Intense Quote": "q",
    "Default Paragraph Font": "span"
}


def slugify(text):
    text = unidecode.unidecode(text).lower()
    r = re.sub(r'[\W_]+', '-', text)
    if r.endswith("-"):
        r = r[:len(r) - 1]
    return r


def htmltotext(htm):
    regex = re.compile(r'<[^>]+>')
    return regex.sub(' ', htm)


def iter_block_items(parent):
    """
    Yield each paragraph and table child within *parent*, in document order.
    Each returned value is an instance of either Table or Paragraph. *parent*
    would most commonly be a reference to a main Document object, but
    also works for a _Cell object, which itself can contain paragraphs and tables.
    """
    if isinstance(parent, doc):
        parent_elm = parent.element.body
    elif isinstance(parent, _Cell):
        parent_elm = parent._tc
    else:
        raise ValueError("something's not right")

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def parsestyle(document, isrun):
    if not isrun:
        css = ""
        font = document.runs[0].font
        if document.paragraph_format.alignment is not None:
            css += "text-align: {0};".format(str(document.paragraph_format.alignment).replace(" (1)", ""))
        if document.paragraph_format.left_indent is not None:
            css += "left: {0};".format(document.paragraph_format.left_indent.pt * 0.1)
        if document.paragraph_format.right_indent is not None:
            css += "right: {0};".format(document.paragraph_format.right_indent.pt * 0.1)
        if document.paragraph_format.line_spacing is not None:
            css += "line-height: {0};".format(document.paragraph_format.line_spacing.pt)
        if font.size is not None:
            css += "font-size: {0};".format(font.size.pt)
        if font.italic is not None and font.italic:
            css += "font-style: italic;"
        if font.bold is not None and font.bold:
            css += "font-weight: bold;"
        if font.underline is not None:
            if font.underline:
                css += "text-decoration-line: underline;"
            else:
                css += "text-decoration-line: " + font.underline + ";"
        if font.highlight_color is not None:
            css += "background_color: #" + str(font.highlight_color) + ";"
        if font.color.rgb is not None:
            css += "color: #" + str(font.color.rgb) + ";"
        return css
    else:
        child_font = document.font
        child_css = ""
        if child_font.size is not None:
            child_css += "font-size: {0};".format(child_font.size.pt)
        if child_font.italic is not None and child_font.italic:
            child_css += "font-style: italic;"
        if child_font.bold is not None and child_font.bold:
            child_css += "font-weight: bold;"
        if child_font.underline is not None:
            if child_font.underline:
                child_css += "text-decoration-line: underline;"
            else:
                child_css += "text-decoration-line: " + str(child_font.underline) + ";"
        if child_font.highlight_color is not None:
            child_css += "background_color: #" + str(child_font.highlight_color) + ";"
        if child_font.color.rgb is not None:
            child_css += "color: #" + str(child_font.color.rgb) + ";"
        return child_css


def generatehtmltag(document):
    global styles
    document.add_run()
    htmlstring = ""
    css = parsestyle(document, False)
    tag = styles[document.style.name]
    htmlstring += "<" + tag + " style='" + css + "'>"
    for run in document.runs:
        child_css = parsestyle(run, True)
        child_tag = styles[run.style.name]
        if run.text != "":
            htmlstring += "<" + child_tag + " style='" + child_css + "'>" + run.text + "</" + child_tag + ">"
    htmlstring += "</" + tag + ">"
    if document.text != "":
        return "{0}".format(htmlstring)
    return ""


def getcontent(file, document):
    global config
    global content
    html = ""
    if not os.path.exists(config["directories"]["output"]):
        os.mkdir(config["directories"]["output"])
    if not os.path.exists(config["directories"]["output"] + "/public"):
        os.mkdir(config["directories"]["output"] + "/public")
    if not os.path.exists(config["directories"]["output"] + "/public/images"):
        os.mkdir(config["directories"]["output"] + "/public/images")
    if config["generator"]["input"] == "docx":
        doc = docx.Document(file)
        doc_properties = doc.core_properties
        html = ""
        images = {}
        id = str(random.randint(10000, 99999))
        imagedir = "/public/images/" + slugify(document["file"]) + id
        if not os.path.exists(config["directories"]["output"] + imagedir):
            os.mkdir(config["directories"]["output"] + imagedir)
        docx2txt.process(file, config["directories"]["output"] + imagedir)
        for r in doc.part.rels.values():
            if isinstance(r._target, docx.parts.image.ImagePart):
                images[r.rId] = os.path.basename(r._target.partname)
        i = 0
        for block in iter_block_items(doc):
            if 'text' in str(block):
                for run in block.runs:
                    xmlstr = str(run.element.xml)
                if 'Graphic' in xmlstr:
                    for rId in images:
                        if rId in xmlstr:
                            htmlstring = "<img class='img-fluid' src='" + imagedir + "/" + images[rId] + "'>"
                            if htmlstring not in html:
                                html += htmlstring
                if block.text is not None:
                    html += generatehtmltag(block)
            elif 'table' in str(block):
                tablehtml = "<table>"
                tab = doc.tables[i]
                for row in tab.rows:
                    tr = "<tr>"
                    for cell in row.cells:
                        tr += "<td>{0}</td>".format(cell.text)
                    tr += "</tr>"
                    tablehtml += tr
                tablehtml += "</table>"
                html += tablehtml
                i += 1
        document["id"] = id
        document["imagedir"] = imagedir
        if doc_properties.created == None:
            document["date"] = datetime.date.today().strftime("%B %d, %Y")
        else:
            document["date"] = doc_properties.created.strftime("%B %d, %Y")
        document["body"] = html
        content[document["file"]] = document


def scancontent():
    """
    Scan Posts
    Document() document.core_properties.created for date
    """
    global config
    global content
    if os.path.exists(config["directories"]["input"]):
        source = Path(config["directories"]["input"] + "/")
        files = source.glob("*")
        for file in files:
            if file.is_dir():
                doctype = file.name
                file = Path(config["directories"]["input"] + "/" + file.name).glob("*." + config["generator"]["input"])
                for filecontent in file:
                    document = {
                        "type": doctype,
                        "file": filecontent.name.split(".")[0],
                        "title": filecontent.name.split(".")[0],
                        "body": "",
                    }
                    if not document["file"] in content:
                        document.update(config["author"])
                        document.update(config["site"])
                        getcontent(filecontent, document)
            else:
                print("Found misplaced file " + file.name + " please categorize your documents correctly. Skipping.")
    json_object = json.dumps(content, indent=4)
    with open('generated_output.json', 'w') as file:
        file.write(json_object)


def parsetemplate(input, type):
    """
    :input: string whom will be added into parsed template
    :type: page, post, category, search, profile
    :rtype: string
    """
    global config
    loadtheme()
    themefile = config['directories']['themes'] + "/" + config['site']['theme'] + "/" + type + ".html"
    if os.path.exists(themefile):
        p = re.compile('(\[\[([a-z]+)\]\])')
        output = str(open(themefile).read())
        matches = p.findall(output)
        for placeholder, token in matches:
            if token in input:
                output = output.replace(placeholder, str(input[token]))
        return output
    else:
        print("Warning!!! Template not found...")


def parsesnippet(input, snippet):
    global theme
    loadtheme()
    p = re.compile('(\[\[([a-z]+)\]\])')
    matches = p.findall(theme["snippets"][snippet])
    output = theme["snippets"][snippet]
    for p, match in matches:
        if match in input:
            output = output.replace(p, str(input[match]))
    return output


def generatehomepage():
    global config
    global content
    homecontent = {}
    homecontent.update(config["author"])
    homecontent.update(config["site"])
    homecontent["body"] = ""
    tempcontent = {}
    date_order = OrderedDict(sorted(content.items()), key=lambda t: t["date"])
    for post in date_order:
        if type(date_order[post]) is dict:
            if content[post]["type"] == "post":
                tempcontent.update(content[post])
                tempcontent["file"] = slugify(tempcontent["file"])
                tempcontent["body"] = htmltotext(tempcontent["body"])
                tempcontent["body"] = (tempcontent["body"][:120] + '..') if len(tempcontent["body"]) > 120 else \
                    tempcontent["body"]
                homecontent["body"] += parsesnippet(tempcontent, "home_post")
                tempcontent = {}
    filename = config["directories"]["output"] + "/index.html"
    outfile = open(filename, "w")
    outfile.write(parsetemplate(homecontent, "home"))
    outfile.close()


def generatehtml():
    print("Started scan.")
    scancontent()
    print("Scan completed. Generating homepage")
    generatehomepage()
    for doc in content:
        document = content[doc]
        if not os.path.exists(config["directories"]["output"] + "/" + document["type"]):
            os.mkdir(config["directories"]["output"] + "/" + document["type"])
        filename = config["directories"]["output"] + "/" + document["type"] + "/" + slugify(document["file"]) + ".html"
        outfile = open(filename, "w")
        outfile.write(parsetemplate(document, document["type"]))
        outfile.close()
    if os.path.exists(config["directories"]["themes"] + "/" + config["site"]["theme"] + "/assets"):
        print("Found theme assets, Copying them.")
        if os.path.exists(config["directories"]["output"] + "/public/assets"):
            shutil.rmtree(config["directories"]["output"] + "/public/assets")
        shutil.copytree(config["directories"]["themes"] + "/" + config["site"]["theme"] + "/assets",
                        config["directories"]["output"] + "/public/assets")
    print("Site generation completed.")


def loadtheme():
    global theme
    global config
    if os.path.exists(config["directories"]["themes"] + "/" + config["site"]["theme"] + "/config.yaml"):
        path = config["directories"]["themes"] + "/" + config["site"]["theme"]
        with open(path + "/config.yaml") as file:
            try:
                theme = yaml.safe_load(file)
            except yaml.YAMLError as exception:
                print(exception)
    else:
        print("Theme configuration file not found!")


def downloadtheme(name):
    global config
    url = "https://api.github.com/repos/devsimsek/WordFlow_themes/tarball/" + name + "_theme"
    try:
        status = urllib.request.urlopen(url)
    except urllib.error.HTTPError:
        print("Theme " + name + " not found.")
        return
    if not os.path.exists(config["directories"]["themes"] + "/" + name):
        if not os.path.exists("temp"):
            os.mkdir("temp")
        if not os.path.exists(config["directories"]["themes"] + "/" + name):
            os.mkdir(config["directories"]["themes"] + "/" + name)
        urllib.request.urlretrieve(url, "temp/" + name + "_theme.tar.gz")
        theme = tarfile.open("temp/" + name + "_theme.tar.gz")
        theme.extractall(config["directories"]["themes"] + "/" + name)
        extractedfile = os.path.commonprefix(theme.getnames())
        theme.close()
        for file in Path(config["directories"]["themes"] + "/" + name + "/" + extractedfile).glob("*"):
            shutil.move(file, config["directories"]["themes"] + "/" + name)
        shutil.rmtree(config["directories"]["themes"] + "/" + name + "/" + extractedfile)
        shutil.rmtree("temp")
        print("Theme installed.")
    else:
        print("Selected theme already exists. Want to reinstall?")
        val = input("(yes, no)> ")
        if val != "yes":
            print("Skipping theme installation.")
        else:
            print("Reinstalling...")
            shutil.rmtree(config["directories"]["themes"] + "/" + name)
            downloadtheme(name)


def clearinstallation():
    global config
    for directory in config["directories"]:
        print("removing " + config["directories"][directory])
        shutil.rmtree(config["directories"][directory])
    if os.path.exists("config.yaml"):
        os.remove("config.yaml")
    if os.path.exists("generated_output.json"):
        os.remove("generated_output.json")


def clearcontent():
    global config
    for directory in config["directories"]:
        if directory == "input":
            continue
        if directory == "themes":
            continue
        print("removing " + config["directories"][directory])
        shutil.rmtree(config["directories"][directory])
        if not os.path.exists(config["directories"][directory]):
            os.mkdir(config["directories"][directory])
        if os.path.exists("generated_output.json"):
            os.remove("generated_output.json")


def initapp():
    """
    Initialize wordflow application
    """
    global config
    print("Welcome to the WordFlow initializer.")
    print("Checking configuration")
    if not os.path.exists("config.yaml"):
        for key in config:
            i = 1
            for opt in config[key]:
                print(
                    "--- " + key.capitalize() + " Configuration (" + str(i) + " of " + str(len(config[key])) + ") ---")
                print("Configuring " + opt + " field.")
                val = input("Value (default: " + config[key][opt] + ")> ")
                if not val == "":
                    config[key][opt] = val
                i += 1
        print("Configuration completed.")
        if not os.path.exists("config.yaml"):
            with open("config.yaml", "w") as file:
                try:
                    yaml.dump(config, file)
                    print("Configuration saved. You can create your documents now.")
                except yaml.YAMLError as exception:
                    print(exception)
        else:
            print("Operation failed. Configuration already exists!")
            val = input(
                "Want to clean install WordFlow? (this will remove every configuration and files.) (yes or no)> ")
            if val != "yes":
                print("Bye :)")
            else:
                clearinstallation()
    else:
        print("Configuration found. Skipping.")
    print("Checking directories")
    for directory in config["directories"]:
        if not os.path.exists(config["directories"][directory]):
            print(config["directories"][directory] + " not exists. Creating.")
            os.mkdir(config["directories"][directory])
        else:
            print(directory + " exists.")
    if os.path.exists(config["directories"]["input"]):
        if not os.path.exists(config["directories"]["input"] + "/post"):
            os.mkdir(config["directories"]["input"] + "/post")
        if not os.path.exists(config["directories"]["input"] + "/page"):
            os.mkdir(config["directories"]["input"] + "/page")
    if config["site"]["theme"] == "default":
        print("Installing Theme")
        downloadtheme(config["site"]["theme"])

    print("Application should be initialized correctly. Thanks for using WordFlow.")
    exit(1)


def argvparser():
    args = sys.argv[1:]
    for arg in args:
        if arg == "init" or arg == "-init":
            initapp()
        elif arg == "generate" or arg == "gen":
            generatehtml()
        elif arg == "clear":
            val = input(
                "Want to clean install WordFlow? (this will remove every configuration and files.) (yes or no)> ")
            if val != "yes":
                print("Bye :)")
            else:
                clearinstallation()
        elif arg == "installtheme" or arg == "theme":
            name = input("Theme name> ")
            downloadtheme(name)
        elif arg == "scan":
            scancontent()
        elif arg == "clearcontent" or arg == "-cc":
            val = input(
                "Want to wipe all generated content? (yes or no)> ")
            if val != "yes":
                print("Bye :)")
            else:
                clearcontent()


def wordflow():
    if not os.path.exists("config.yaml"):
        print("Warning: Configuration file not found. Launching initializer.")
        initapp()
    else:
        global config
        with open("config.yaml") as file:
            try:
                config = yaml.safe_load(file)
            except yaml.YAMLError as exception:
                print(exception)


if __name__ == "__main__":
    wordflow()
    argvparser()
else:
    print("Illegal Launch Option")
